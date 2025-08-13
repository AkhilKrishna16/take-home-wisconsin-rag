#!/usr/bin/env python3
"""
Document Processor for Legal Documents

This module provides comprehensive document processing capabilities for legal documents,
including text extraction, metadata analysis, and intelligent chunking for RAG systems.
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
from datetime import datetime
import hashlib

# Import the document chunker and vector database
from .document_chunker import DocumentChunker
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from vector_db.vector_database import LegalVectorDatabase

# PDF Processing
try:
    from PyPDF2 import PdfReader
    from pdfminer.high_level import extract_text
    from pdfminer.layout import LAParams
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PDF processing libraries not available. Install with: pip install PyPDF2 pdfminer.six")

# DOCX Processing
try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("DOCX processing library not available. Install with: pip install python-docx")

# Text Processing
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.stem import WordNetLemmatizer
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False
    print("NLTK not available. Install with: pip install nltk")

# Additional libraries
import requests
from urllib.parse import urlparse
import mimetypes

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Document processor with singleton pattern to prevent multiple initializations.
    """
    
    _instance = None
    _initialized = False
    
    def __new__(cls, *args, **kwargs):
        """Singleton pattern to ensure only one instance exists."""
        if cls._instance is None:
            cls._instance = super(DocumentProcessor, cls).__new__(cls)
        return cls._instance
    
    def __init__(self, output_dir: str = "processed_documents", chunk_size: int = 1000, chunk_overlap: int = 200, use_vector_db: bool = True):
        # Only initialize once
        if self._initialized:
            return
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Initialize document chunker
        self.chunker = DocumentChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        
        # Initialize vector database if requested
        self.vector_db = None
        if use_vector_db:
            try:
                self.vector_db = LegalVectorDatabase()
                logger.info("Vector database initialized successfully")
            except Exception as e:
                logger.warning(f"Could not initialize vector database: {e}")
                self.vector_db = None
        
        # Initialize NLTK components if available
        self.nltk_available = NLTK_AVAILABLE
        if self.nltk_available:
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
                nltk.download('wordnet', quiet=True)
                nltk.download('averaged_perceptron_tagger', quiet=True)
                self.lemmatizer = WordNetLemmatizer()
                self.stop_words = set(stopwords.words('english'))
            except Exception as e:
                logger.warning(f"Could not initialize NLTK components: {e}")
                self.nltk_available = False
        
        # Document type patterns
        self.document_patterns = {
            'case_law': {
                'keywords': ['case', 'court', 'judgment', 'opinion', 'appeal', 'petitioner', 'respondent'],
                'patterns': [
                    r'Case No\.?\s*[A-Z0-9\-]+',
                    r'In the [A-Z\s]+ Court',
                    r'Opinion of the Court',
                    r'Filed\s+\d{1,2}/\d{1,2}/\d{4}'
                ]
            },
            'policy': {
                'keywords': ['policy', 'procedure', 'guideline', 'regulation', 'standard', 'protocol'],
                'patterns': [
                    r'Policy\s+No\.?\s*[A-Z0-9\-]+',
                    r'Effective Date:\s*\d{1,2}/\d{1,2}/\d{4}',
                    r'Department\s+Policy',
                    r'Standard Operating Procedure'
                ]
            },
            'training': {
                'keywords': ['training', 'course', 'module', 'lesson', 'instruction', 'education'],
                'patterns': [
                    r'Training\s+Module\s*\d+',
                    r'Course\s+Objective',
                    r'Learning\s+Outcome',
                    r'Training\s+Material'
                ]
            }
        }
        
        # Mark as initialized
        self._initialized = True
    
    def process_document(self, file_path: str, document_type: Optional[str] = None) -> Dict[str, Any]:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Determine file type
        file_type = self._get_file_type(file_path)
        
        # Extract text content
        text_content = self._extract_text(file_path, file_type)
        
        # Auto-detect document type if not provided
        if not document_type:
            document_type = self._detect_document_type(text_content)
        
        # Create intelligent chunks
        chunks = self.chunker.chunk_document(text_content, document_type)
        
        # Process based on document type
        processed_data = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_type': file_type,
            'document_type': document_type,
            'processed_at': datetime.now().isoformat(),
            'file_hash': self._calculate_file_hash(file_path),
            'metadata': self._extract_metadata(file_path, file_type),
            'content': text_content,
            'summary': self._generate_summary(text_content, document_type),
            'key_entities': self._extract_key_entities(text_content, document_type),
            'structure': self._extract_structure(text_content, document_type),
            'tags': self._generate_tags(text_content, document_type),
            'chunks': chunks,
            'chunk_count': len(chunks)
        }
        
        # Save processed document
        self._save_processed_document(processed_data)
        
        # Index in vector database if available
        if self.vector_db and chunks:
            document_id = f"{file_path.stem}_{processed_data['file_hash'][:8]}"
            
            # Add file name to each chunk's metadata
            for chunk in chunks:
                if 'metadata' not in chunk:
                    chunk['metadata'] = {}
                chunk['metadata']['file_name'] = file_path.name
                chunk['metadata']['original_file_name'] = file_path.name
            
            success = self.vector_db.index_document_chunks(chunks, document_id)
            if success:
                processed_data['vector_db_indexed'] = True
                processed_data['vector_db_id'] = document_id
            else:
                processed_data['vector_db_indexed'] = False
        
        return processed_data
    
    def _get_file_type(self, file_path: Path) -> str:
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        if file_path.suffix.lower() == '.pdf':
            return 'pdf'
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return 'docx'
        elif file_path.suffix.lower() in ['.txt', '.md', '.rtf']:
            return 'text'
        elif file_path.suffix.lower() in ['.csv', '.tsv']:
            return 'spreadsheet'
        elif file_path.suffix.lower() in ['.html', '.htm']:
            return 'html'
        else:
            return 'unknown'
    
    def _extract_text(self, file_path: Path, file_type: str) -> str:
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'text':
                return self._extract_text_file(file_path)
            elif file_type == 'spreadsheet':
                return self._extract_spreadsheet_text(file_path)
            elif file_type == 'html':
                return self._extract_html_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        try:
            # Try pdfminer first
            text = extract_text(str(file_path), laparams=LAParams())
            if text.strip():
                return text
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                # If still no text, try OCR with pytesseract
                if not text.strip():
                    try:
                        import pytesseract
                        from PIL import Image
                        import fitz  # PyMuPDF
                        
                        # Open PDF with PyMuPDF
                        doc = fitz.open(str(file_path))
                        text = ""
                        
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            # Get page as image
                            pix = page.get_pixmap()
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            
                            # Extract text using OCR
                            page_text = pytesseract.image_to_string(img)
                            text += page_text + "\n"
                        
                        doc.close()
                        logger.info(f"Used OCR to extract text from {file_path}")
                        return text
                        
                    except ImportError:
                        logger.warning("OCR libraries not available. Install with: pip install pytesseract Pillow PyMuPDF")
                    except Exception as ocr_error:
                        logger.error(f"OCR extraction failed: {ocr_error}")
                
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not available")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_text_file(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file with any encoding: {file_path}")
    
    def _extract_spreadsheet_text(self, file_path: Path) -> str:
        try:
            import csv
            text = ""
            with open(file_path, 'r', encoding='utf-8') as file:
                if file_path.suffix.lower() == '.csv':
                    reader = csv.reader(file)
                else:  # .tsv
                    reader = csv.reader(file, delimiter='\t')
                
                for row in reader:
                    text += ' | '.join(row) + '\n'
            return text
        except Exception as e:
            logger.error(f"Error extracting spreadsheet text: {e}")
            return ""
    
    def _extract_html_text(self, file_path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                # Get text and clean it up
                text = soup.get_text()
                # Remove extra whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                return text
        except ImportError:
            logger.warning("BeautifulSoup not available for HTML processing")
            # Fallback to basic text extraction
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting HTML text: {e}")
            return ""
    
    def _detect_document_type(self, text_content: str) -> str:
        text_lower = text_content.lower()
        scores = {}
        
        for doc_type, patterns in self.document_patterns.items():
            score = 0

            for keyword in patterns['keywords']:
                if keyword in text_lower:
                    score += 1

            for pattern in patterns['patterns']:
                if re.search(pattern, text_content, re.IGNORECASE):
                    score += 2
            
            scores[doc_type] = score

        if scores:
            return max(scores, key=scores.get)
        else:
            return 'general'
    
    def _extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        metadata = {
            'file_size': file_path.stat().st_size,
            'created_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            'modified_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }
        
        if file_type == 'pdf' and PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    if reader.metadata:
                        metadata['pdf_metadata'] = dict(reader.metadata)
                    metadata['page_count'] = len(reader.pages)
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        elif file_type == 'docx' and DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                metadata['paragraph_count'] = len(doc.paragraphs)
                metadata['table_count'] = len(doc.tables)
            except Exception as e:
                logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return metadata
    
    def _generate_summary(self, text_content: str, document_type: str) -> str:
        if not text_content.strip():
            return "No content available for summary."
        
        sentences = sent_tokenize(text_content) if self.nltk_available else text_content.split('.')
        
        if document_type == 'case_law':
            summary_sentences = sentences[:3]  # First 3 sentences
        elif document_type == 'policy':
            summary_sentences = sentences[:2]  # First 2 sentences
        elif document_type == 'training':
            summary_sentences = sentences[:4]  # First 4 sentences
        else:
            summary_sentences = sentences[:3]
        
        return ' '.join(summary_sentences)
    
    def _extract_key_entities(self, text_content: str, document_type: str) -> Dict[str, List[str]]:
        entities = {
            'dates': [],
            'names': [],
            'organizations': [],
            'locations': [],
            'references': []
        }
        
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
        ]
        
        for pattern in date_patterns:
            dates = re.findall(pattern, text_content, re.IGNORECASE)
            entities['dates'].extend(dates)
        
        if document_type == 'case_law':
            case_patterns = [
                r'Case\s+No\.?\s*[A-Z0-9\-]+',
                r'Docket\s+No\.?\s*[A-Z0-9\-]+',
                r'Citation:\s*[A-Z0-9\s]+'
            ]
            for pattern in case_patterns:
                refs = re.findall(pattern, text_content, re.IGNORECASE)
                entities['references'].extend(refs)
        
        elif document_type == 'policy':
            policy_patterns = [
                r'Policy\s+No\.?\s*[A-Z0-9\-]+',
                r'Procedure\s+No\.?\s*[A-Z0-9\-]+'
            ]
            for pattern in policy_patterns:
                refs = re.findall(pattern, text_content, re.IGNORECASE)
                entities['references'].extend(refs)
        
        return entities
    
    def _extract_structure(self, text_content: str, document_type: str) -> Dict[str, Any]:
        structure = {
            'sections': [],
            'headings': [],
            'paragraphs': len(text_content.split('\n\n')),
            'word_count': len(text_content.split())
        }

        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and not line.endswith('.'):
                if re.match(r'^[A-Z][A-Z\s\d]+$', line) or line.isupper():
                    structure['headings'].append(line)
        
        return structure
    
    def _generate_tags(self, text_content: str, document_type: str) -> List[str]:
        tags = [document_type]
        
        if document_type == 'case_law':
            tags.extend(['legal', 'court', 'judgment'])
        elif document_type == 'policy':
            tags.extend(['policy', 'procedure', 'guideline'])
        elif document_type == 'training':
            tags.extend(['training', 'education', 'learning'])
        
        text_lower = text_content.lower()
        
        if any(word in text_lower for word in ['confidential', 'secret', 'private']):
            tags.append('confidential')
        
        if any(word in text_lower for word in ['urgent', 'immediate', 'emergency']):
            tags.append('urgent')
        
        return list(set(tags))
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _save_processed_document(self, processed_data: Dict[str, Any]) -> None:
        file_name = Path(processed_data['file_name']).stem
        output_file = self.output_dir / f"{file_name}_processed.json"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(processed_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Processed document saved to: {output_file}")
    
    def process_directory(self, directory_path: str, file_types: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if file_types is None:
            file_types = ['.pdf', '.docx', '.doc', '.txt', '.md']
        
        processed_documents = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                try:
                    logger.info(f"Processing: {file_path}")
                    processed_data = self.process_document(str(file_path))
                    processed_documents.append(processed_data)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
        
        return processed_documents
    
    def get_processing_stats(self) -> Dict[str, Any]:
        processed_files = list(self.output_dir.glob('*_processed.json'))
        
        stats = {
            'total_processed': len(processed_files),
            'document_types': {},
            'file_types': {},
            'total_size': 0
        }
        
        for file_path in processed_files:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                doc_type = data.get('document_type', 'unknown')
                stats['document_types'][doc_type] = stats['document_types'].get(doc_type, 0) + 1
                
                file_type = data.get('file_type', 'unknown')
                stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
                
                stats['total_size'] += data.get('metadata', {}).get('file_size', 0)
                
            except Exception as e:
                logger.error(f"Error reading processed file {file_path}: {e}")
        
        return stats
    
    def get_chunks_for_rag(self, file_path: str, document_type: Optional[str] = None) -> List[Dict[str, Any]]:
        """Get document chunks optimized for RAG systems."""
        processed_data = self.process_document(file_path, document_type)
        return processed_data.get('chunks', [])
    
    def get_chunks_by_metadata(self, chunks: List[Dict[str, Any]], metadata_filter: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Filter chunks by metadata (e.g., statute numbers, case citations)."""
        filtered_chunks = []
        
        for chunk in chunks:
            chunk_metadata = chunk.get('metadata', {})
            matches_filter = True
            
            for key, value in metadata_filter.items():
                if key in chunk_metadata:
                    if isinstance(value, list):
                        # Check if any value in the filter matches any value in chunk metadata
                        if not any(v in chunk_metadata[key] for v in value):
                            matches_filter = False
                            break
                    else:
                        # Direct value comparison
                        if value not in chunk_metadata[key]:
                            matches_filter = False
                            break
                else:
                    matches_filter = False
                    break
            
            if matches_filter:
                filtered_chunks.append(chunk)
        
        return filtered_chunks
    
    def search_documents(self, query: str, top_k: int = 10, filter_metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Search documents using vector database."""
        if not self.vector_db:
            raise RuntimeError("Vector database not initialized")
        
        return self.vector_db.search_legal_documents(query, top_k, filter_metadata)
    
    def search_by_statute(self, statute_number: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search for documents containing specific statute references."""
        if not self.vector_db:
            raise RuntimeError("Vector database not initialized")
        
        return self.vector_db.search_by_statute(statute_number, top_k)
    
    def search_by_case_citation(self, case_name: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search for documents containing specific case citations."""
        if not self.vector_db:
            raise RuntimeError("Vector database not initialized")
        
        return self.vector_db.search_by_case_citation(case_name, top_k)
    
    def search_by_document_type(self, document_type: str, query: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search within a specific document type."""
        if not self.vector_db:
            raise RuntimeError("Vector database not initialized")
        
        return self.vector_db.search_by_document_type(document_type, query, top_k)
    
    def get_vector_db_stats(self) -> Dict[str, Any]:
        """Get vector database statistics."""
        if not self.vector_db:
            raise RuntimeError("Vector database not initialized")
        
        return self.vector_db.get_index_stats()


if __name__ == "__main__":
    processor = DocumentProcessor()
    
    print("Document Processor initialized successfully!")
    print("Available methods:")
    print("- process_document(file_path, document_type=None)")
    print("- process_directory(directory_path, file_types=None)")
    print("- get_processing_stats()")
    print("- get_chunks_for_rag(file_path, document_type=None)")
    print("- get_chunks_by_metadata(chunks, metadata_filter)")
    print("- search_documents(query, top_k=10, filter_metadata=None)")
    print("- search_by_statute(statute_number, top_k=10)")
    print("- search_by_case_citation(case_name, top_k=10)")
    print("- search_by_document_type(document_type, query, top_k=10)")
    print("- get_vector_db_stats()")
