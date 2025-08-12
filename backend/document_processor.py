import os
import re
import json
import logging
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
from datetime import datetime
import hashlib

# PDF Processing
try:
    from PyPDF2 import PdfReader
    from pdfminer.high_level import extract_text
    from pdfminer.layout import LAParams
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("PDF processing libraries not available. Install with: pip install PyPDF2 pdfminer.six")

# DOCX Processing
try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("DOCX processing library not available. Install with: pip install python-docx")

# Text Processing
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.stem import WordNetLemmatizer
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False
    print("NLTK not available. Install with: pip install nltk")

# Additional libraries
import requests
from urllib.parse import urlparse
import mimetypes

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    A comprehensive document processor for legal and policy documents.
    Supports PDF, DOCX, and various text formats.
    """
    
    def __init__(self, output_dir: str = "processed_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Initialize NLTK components if available
        if NLTK_AVAILABLE:
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
                nltk.download('wordnet', quiet=True)
                nltk.download('averaged_perceptron_tagger', quiet=True)
                self.lemmatizer = WordNetLemmatizer()
                self.stop_words = set(stopwords.words('english'))
            except Exception as e:
                logger.warning(f"Could not initialize NLTK components: {e}")
                NLTK_AVAILABLE = False
        
        # Document type patterns
        self.document_patterns = {
            'case_law': {
                'keywords': ['case', 'court', 'judgment', 'opinion', 'appeal', 'petitioner', 'respondent'],
                'patterns': [
                    r'Case No\.?\s*[A-Z0-9\-]+',
                    r'In the [A-Z\s]+ Court',
                    r'Opinion of the Court',
                    r'Filed\s+\d{1,2}/\d{1,2}/\d{4}'
                ]
            },
            'policy': {
                'keywords': ['policy', 'procedure', 'guideline', 'regulation', 'standard', 'protocol'],
                'patterns': [
                    r'Policy\s+No\.?\s*[A-Z0-9\-]+',
                    r'Effective Date:\s*\d{1,2}/\d{1,2}/\d{4}',
                    r'Department\s+Policy',
                    r'Standard Operating Procedure'
                ]
            },
            'training': {
                'keywords': ['training', 'course', 'module', 'lesson', 'instruction', 'education'],
                'patterns': [
                    r'Training\s+Module\s*\d+',
                    r'Course\s+Objective',
                    r'Learning\s+Outcome',
                    r'Training\s+Material'
                ]
            }
        }
    
    def process_document(self, file_path: str, document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a document and extract structured information.
        
        Args:
            file_path: Path to the document
            document_type: Optional document type (case_law, policy, training)
            
        Returns:
            Dictionary containing processed document information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Determine file type
        file_type = self._get_file_type(file_path)
        
        # Extract text content
        text_content = self._extract_text(file_path, file_type)
        
        # Auto-detect document type if not provided
        if not document_type:
            document_type = self._detect_document_type(text_content)
        
        # Process based on document type
        processed_data = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_type': file_type,
            'document_type': document_type,
            'processed_at': datetime.now().isoformat(),
            'file_hash': self._calculate_file_hash(file_path),
            'metadata': self._extract_metadata(file_path, file_type),
            'content': text_content,
            'summary': self._generate_summary(text_content, document_type),
            'key_entities': self._extract_key_entities(text_content, document_type),
            'structure': self._extract_structure(text_content, document_type),
            'tags': self._generate_tags(text_content, document_type)
        }
        
        # Save processed document
        self._save_processed_document(processed_data)
        
        return processed_data
    
    def _get_file_type(self, file_path: Path) -> str:
        """Determine the file type based on extension and content."""
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        if file_path.suffix.lower() == '.pdf':
            return 'pdf'
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return 'docx'
        elif file_path.suffix.lower() in ['.txt', '.md']:
            return 'text'
        else:
            return 'unknown'
    
    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text content from different file types."""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'text':
                return self._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        try:
            # Try pdfminer first for better text extraction
            text = extract_text(str(file_path), laparams=LAParams())
            if text.strip():
                return text
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not available")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file with any encoding: {file_path}")
    
    def _detect_document_type(self, text_content: str) -> str:
        """Auto-detect document type based on content analysis."""
        text_lower = text_content.lower()
        scores = {}
        
        for doc_type, patterns in self.document_patterns.items():
            score = 0

            for keyword in patterns['keywords']:
                if keyword in text_lower:
                    score += 1

            for pattern in patterns['patterns']:
                if re.search(pattern, text_content, re.IGNORECASE):
                    score += 2
            
            scores[doc_type] = score

        if scores:
            return max(scores, key=scores.get)
        else:
            return 'general'
    
    def _extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """Extract metadata from the document."""
        metadata = {
            'file_size': file_path.stat().st_size,
            'created_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            'modified_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
        }
        
        if file_type == 'pdf' and PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    if reader.metadata:
                        metadata['pdf_metadata'] = dict(reader.metadata)
                    metadata['page_count'] = len(reader.pages)
            except Exception as e:
                logger.warning(f"Could not extract PDF metadata: {e}")
        
        elif file_type == 'docx' and DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                metadata['paragraph_count'] = len(doc.paragraphs)
                metadata['table_count'] = len(doc.tables)
            except Exception as e:
                logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return metadata
    
    def _generate_summary(self, text_content: str, document_type: str) -> str:
        """Generate a summary of the document content."""
        if not text_content.strip():
            return "No content available for summary."
        
        # Simple summary based on document type
        sentences = sent_tokenize(text_content) if NLTK_AVAILABLE else text_content.split('.')
        
        if document_type == 'case_law':
            summary_sentences = sentences[:3]  # First 3 sentences
        elif document_type == 'policy':
            # Focus on policy scope and key points
            summary_sentences = sentences[:2]  # First 2 sentences
        elif document_type == 'training':
            # Focus on learning objectives
            summary_sentences = sentences[:4]  # First 4 sentences
        else:
            summary_sentences = sentences[:3]
        
        return ' '.join(summary_sentences)
    
    def _extract_key_entities(self, text_content: str, document_type: str) -> Dict[str, List[str]]:
        """Extract key entities from the document."""
        entities = {
            'dates': [],
            'names': [],
            'organizations': [],
            'locations': [],
            'references': []
        }
        
        # Extract dates
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
        ]
        
        for pattern in date_patterns:
            dates = re.findall(pattern, text_content, re.IGNORECASE)
            entities['dates'].extend(dates)
        
        # Extract case numbers and references
        if document_type == 'case_law':
            case_patterns = [
                r'Case\s+No\.?\s*[A-Z0-9\-]+',
                r'Docket\s+No\.?\s*[A-Z0-9\-]+',
                r'Citation:\s*[A-Z0-9\s]+'
            ]
            for pattern in case_patterns:
                refs = re.findall(pattern, text_content, re.IGNORECASE)
                entities['references'].extend(refs)
        
        # Extract policy numbers
        elif document_type == 'policy':
            policy_patterns = [
                r'Policy\s+No\.?\s*[A-Z0-9\-]+',
                r'Procedure\s+No\.?\s*[A-Z0-9\-]+'
            ]
            for pattern in policy_patterns:
                refs = re.findall(pattern, text_content, re.IGNORECASE)
                entities['references'].extend(refs)
        
        return entities
    
    def _extract_structure(self, text_content: str, document_type: str) -> Dict[str, Any]:
        """Extract document structure and sections."""
        structure = {
            'sections': [],
            'headings': [],
            'paragraphs': len(text_content.split('\n\n')),
            'word_count': len(text_content.split())
        }
        
        # Extract headings (lines that are shorter and end with common patterns)
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 100 and not line.endswith('.'):
                if re.match(r'^[A-Z][A-Z\s\d]+$', line) or line.isupper():
                    structure['headings'].append(line)
        
        return structure
    
    def _generate_tags(self, text_content: str, document_type: str) -> List[str]:
        """Generate tags for the document."""
        tags = [document_type]
        
        # Add document type specific tags
        if document_type == 'case_law':
            tags.extend(['legal', 'court', 'judgment'])
        elif document_type == 'policy':
            tags.extend(['policy', 'procedure', 'guideline'])
        elif document_type == 'training':
            tags.extend(['training', 'education', 'learning'])
        
        # Add content-based tags
        text_lower = text_content.lower()
        
        if any(word in text_lower for word in ['confidential', 'secret', 'private']):
            tags.append('confidential')
        
        if any(word in text_lower for word in ['urgent', 'immediate', 'emergency']):
            tags.append('urgent')
        
        return list(set(tags))
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of the file."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _save_processed_document(self, processed_data: Dict[str, Any]) -> None:
        """Save processed document data to JSON file."""
        file_name = Path(processed_data['file_name']).stem
        output_file = self.output_dir / f"{file_name}_processed.json"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(processed_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Processed document saved to: {output_file}")
    
    def process_directory(self, directory_path: str, file_types: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Process all documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            file_types: Optional list of file extensions to process (e.g., ['.pdf', '.docx'])
            
        Returns:
            List of processed document data
        """
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if file_types is None:
            file_types = ['.pdf', '.docx', '.doc', '.txt', '.md']
        
        processed_documents = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                try:
                    logger.info(f"Processing: {file_path}")
                    processed_data = self.process_document(str(file_path))
                    processed_documents.append(processed_data)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
        
        return processed_documents
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get statistics about processed documents."""
        processed_files = list(self.output_dir.glob('*_processed.json'))
        
        stats = {
            'total_processed': len(processed_files),
            'document_types': {},
            'file_types': {},
            'total_size': 0
        }
        
        for file_path in processed_files:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Count document types
                doc_type = data.get('document_type', 'unknown')
                stats['document_types'][doc_type] = stats['document_types'].get(doc_type, 0) + 1
                
                # Count file types
                file_type = data.get('file_type', 'unknown')
                stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
                
                # Add file size
                stats['total_size'] += data.get('metadata', {}).get('file_size', 0)
                
            except Exception as e:
                logger.error(f"Error reading processed file {file_path}: {e}")
        
        return stats


# Example usage and testing
if __name__ == "__main__":
    # Initialize processor
    processor = DocumentProcessor()
    
    # Example: Process a single document
    # processed = processor.process_document("path/to/document.pdf")
    # print(json.dumps(processed, indent=2))
    
    # Example: Process a directory
    # processed_docs = processor.process_directory("path/to/documents/")
    # print(f"Processed {len(processed_docs)} documents")
    
    # Example: Get processing stats
    # stats = processor.get_processing_stats()
    # print(json.dumps(stats, indent=2))
    
    print("Document Processor initialized successfully!")
    print("Available methods:")
    print("- process_document(file_path, document_type=None)")
    print("- process_directory(directory_path, file_types=None)")
    print("- get_processing_stats()")
